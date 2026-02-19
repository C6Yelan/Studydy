"""DOCX extractor implementation."""

from __future__ import annotations

from pathlib import Path
from uuid import uuid4

from docx import Document

from app.services.extraction.base import BaseExtractor
from app.services.extraction.models import ExtractionWarning, Segment


class DocxExtractor(BaseExtractor):
    def extract(
        self,
        file_path: Path,
        source_path: str,
        doc_uid: str,
        extracted_at: str,
    ) -> tuple[list[Segment], list[ExtractionWarning]]:
        document = Document(str(file_path))
        segments: list[Segment] = []
        warnings: list[ExtractionWarning] = []

        for paragraph_index, paragraph in enumerate(document.paragraphs, start=1):
            text = paragraph.text or ""
            locator = {"paragraph": paragraph_index}
            segments.append(
                Segment(
                    segment_id=str(uuid4()),
                    doc_uid=doc_uid,
                    source_path=source_path,
                    file_type="docx",
                    unit_type="paragraph",
                    unit_index=paragraph_index,
                    text=text,
                    locator=locator,
                    extracted_at=extracted_at,
                )
            )
            if not text.strip():
                warnings.append(
                    ExtractionWarning(
                        code="empty_text_paragraph",
                        message="Paragraph contains no extractable text.",
                        locator=locator,
                    )
                )

        if not segments:
            warnings.append(
                ExtractionWarning(
                    code="docx_no_paragraphs",
                    message="DOCX contains no paragraphs.",
                )
            )

        if document.tables:
            warnings.append(
                ExtractionWarning(
                    code="docx_tables_ignored",
                    message=(
                        "DOCX table cells are not extracted in T2 MVP; only paragraphs are included."
                    ),
                )
            )

        return segments, warnings
